import os
import sys
import re

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx 库: pip install python-docx")
    sys.exit(1)

def parse_markdown_to_docx(md_path, docx_path):
    doc = docx.Document()

    # 彻底解决中文字体的终极魔法
    # 1. 拦截底层 XML 的默认字体设定
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    # 2. 暴力覆盖所有样式的基础字体
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = u'微软雅黑'
            if hasattr(style.element, 'rPr') and style.element.rPr is not None:
                style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    if not os.path.exists(md_path):
        print(f"找不到文件: {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line == '[PAGE_BREAK]':
            doc.add_page_break()
            continue

        if line.startswith('# '):
            h = doc.add_heading(line[2:], 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], 2)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 112, 192)
                else:
                    run = p.add_run(part)

                # 对每一个 run 再次进行地毯式强制覆盖
                run.font.name = u'微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    # 最后一道防线：直接修改底层 XML 的默认 document default
    for font in doc.element.xpath('//w:rFonts'):
        font.set(qn('w:ascii'), u'微软雅黑')
        font.set(qn('w:eastAsia'), u'微软雅黑')
        font.set(qn('w:hAnsi'), u'微软雅黑')
        font.set(qn('w:cs'), u'微软雅黑')

    doc.save(docx_path)
    print(f"✅ 沉浸式 Word 文档已生成，全篇字体已锁定: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python export-docx.py <输入.md> <输出.docx>")
        sys.exit(1)
    parse_markdown_to_docx(sys.argv[1], sys.argv[2])
